from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required, user_passes_test
from django.http import JsonResponse
from django.core.mail import send_mail
from django.conf import settings
from django.template.loader import render_to_string
from django.utils import timezone
from .models import FieldJob, FieldJobAssignment, FieldReport
from .forms import FieldJobForm, FieldReportForm
from core.models import User
import json
from django.contrib import messages
from django.core.exceptions import PermissionDenied
from django.db.models import Q
from django.views.decorators.http import require_GET
import logging

logger = logging.getLogger(__name__)


def is_admin(user):
    return user.is_authenticated and user.role == 'admin'


def is_staff(user):
    return user.is_authenticated and user.role == 'staff'


@login_required
@user_passes_test(is_admin)
def field_job_list(request):
    field_jobs = FieldJob.objects.all().order_by('-created_at')
    return render(request, 'fieldwork/field_job_list.html', {'field_jobs': field_jobs})


@login_required
@user_passes_test(is_admin)
def field_job_create(request):
    """
    Create a new field job with staff assignments and send notifications
    """
    if not request.user.role == 'admin':
        raise PermissionDenied("You don't have permission to create field jobs.")

    if request.method == 'POST':
        form = FieldJobForm(request.POST)
        if form.is_valid():
            try:
                # Start transaction for data consistency
                field_job = form.save(commit=False)
                field_job.created_by = request.user

                # Ensure the datetime is properly saved
                field_job.save()

                # Create assignments for selected staff
                staff_members = form.cleaned_data['staff_members']
                assignment_count = 0
                email_errors = []

                for staff in staff_members:
                    try:
                        # Create assignment
                        assignment = FieldJobAssignment.objects.create(
                            field_job=field_job,
                            staff=staff
                        )
                        assignment_count += 1

                        # Send email to POC with approval code
                        try:
                            send_poc_approval_email(field_job, assignment)
                        except Exception as e:
                            logger.error(f"Failed to send POC email for assignment {assignment.id}: {str(e)}")
                            email_errors.append(f"POC email for {staff.get_full_name()}")

                        # Send email to staff with job information
                        try:
                            send_staff_assignment_email(field_job, staff)
                        except Exception as e:
                            logger.error(f"Failed to send staff email to {staff.email}: {str(e)}")
                            email_errors.append(f"Staff email to {staff.get_full_name()}")

                    except Exception as e:
                        logger.error(f"Failed to create assignment for staff {staff.id}: {str(e)}")
                        messages.warning(request, f"Failed to assign {staff.get_full_name()} to the job.")

                # Prepare success message
                if assignment_count > 0:
                    success_message = f"Field job for '{field_job.client_name}' created successfully! "
                    success_message += f"{assignment_count} staff member(s) assigned."

                    if email_errors:
                        success_message += f" Note: Failed to send emails for: {', '.join(email_errors)}"

                    messages.success(request, success_message)
                else:
                    messages.error(request,
                                   "Field job created but no staff members were assigned. Please edit the job to assign staff.")

                return redirect('fieldwork:field_job_list')

            except Exception as e:
                logger.error(f"Error creating field job: {str(e)}")
                messages.error(request, f"An error occurred while creating the field job: {str(e)}")
                # Form will be re-rendered with existing data
        else:
            # Form is invalid
            error_count = len(form.errors)
            if error_count == 1:
                messages.error(request, "Please correct the error below.")
            else:
                messages.error(request, f"Please correct the {error_count} errors below.")

            # Log form errors for debugging
            logger.warning(f"Field job form validation failed: {form.errors}")

    else:
        # GET request - initialize form
        form = FieldJobForm()

    # Prepare context with additional info
    context = {
        'form': form,
        'title': 'Create Field Job',
        'staff_count': User.objects.filter(role='staff', is_active=True, is_approved=True).count()
    }

    return render(request, 'fieldwork/field_job_form.html', context)


@login_required
@user_passes_test(is_admin)
def field_job_edit(request, pk):
    """
    Edit an existing field job - send emails to ALL assigned staff and POC for updates
    """
    field_job = get_object_or_404(FieldJob, pk=pk)

    if not request.user.role == 'admin':
        raise PermissionDenied("You don't have permission to edit field jobs.")

    if request.method == 'POST':
        form = FieldJobForm(request.POST, instance=field_job)
        if form.is_valid():
            try:
                # Save the main field job
                field_job = form.save()

                # Get current and new staff assignments
                current_assignments = FieldJobAssignment.objects.filter(field_job=field_job)
                current_staff_ids = set(current_assignments.values_list('staff_id', flat=True))
                new_staff_ids = set(staff.id for staff in form.cleaned_data['staff_members'])

                # Find staff to add and remove
                staff_to_add_ids = new_staff_ids - current_staff_ids
                staff_to_remove_ids = current_staff_ids - new_staff_ids

                # Remove assignments for staff that were unselected
                if staff_to_remove_ids:
                    assignments_removed = FieldJobAssignment.objects.filter(
                        field_job=field_job,
                        staff_id__in=staff_to_remove_ids
                    )
                    removed_count = assignments_removed.count()
                    assignments_removed.delete()
                else:
                    removed_count = 0

                # Add new assignments for staff that were added
                assignment_count = 0
                for staff_id in staff_to_add_ids:
                    try:
                        staff = User.objects.get(id=staff_id, role='staff', is_active=True, is_approved=True)

                        # Create new assignment
                        assignment = FieldJobAssignment.objects.create(
                            field_job=field_job,
                            staff=staff
                        )
                        assignment_count += 1
                    except User.DoesNotExist:
                        logger.error(f"Staff member with ID {staff_id} not found or not eligible")
                        messages.warning(request,
                                         f"Staff member (ID: {staff_id}) not found or not eligible for assignment.")
                    except Exception as e:
                        logger.error(f"Failed to create assignment for staff {staff_id}: {str(e)}")
                        messages.warning(request, f"Failed to assign staff member (ID: {staff_id}).")

                # Send emails to ALL currently assigned staff and POC for updates
                email_errors = []
                updated_assignments = FieldJobAssignment.objects.filter(field_job=field_job)

                # Send update emails to all assigned staff
                for assignment in updated_assignments:
                    try:
                        send_staff_update_email(field_job, assignment.staff)
                    except Exception as e:
                        logger.error(f"Failed to send update email to staff {assignment.staff.email}: {str(e)}")
                        email_errors.append(f"Staff update email to {assignment.staff.get_full_name()}")

                # Send update emails to POC for all assignments
                for assignment in updated_assignments:
                    try:
                        send_poc_update_email(field_job, assignment)
                    except Exception as e:
                        logger.error(f"Failed to send POC update email for assignment {assignment.id}: {str(e)}")
                        email_errors.append(f"POC update email for {assignment.staff.get_full_name()}")

                # Prepare success message
                success_message = f"Field job for '{field_job.client_name}' updated successfully!"

                changes = []
                if assignment_count > 0:
                    changes.append(f"added {assignment_count} new staff")
                if removed_count > 0:
                    changes.append(f"removed {removed_count} staff")

                if changes:
                    success_message += f" Changes: {', '.join(changes)}."

                # Add email notification info
                email_count = updated_assignments.count()
                if email_count > 0:
                    success_message += f" Update notifications sent to {email_count} staff member(s) and POC."

                if email_errors:
                    success_message += f" Email issues: {', '.join(email_errors)}"

                # If no changes to staff, mention that updates were still sent
                if assignment_count == 0 and removed_count == 0:
                    success_message += " Update notifications sent to all assigned staff and POC."

                messages.success(request, success_message)
                return redirect('fieldwork:field_job_list')

            except Exception as e:
                logger.error(f"Error updating field job {pk}: {str(e)}")
                messages.error(request, f"An error occurred while updating the field job: {str(e)}")
        else:
            error_count = len(form.errors)
            if error_count == 1:
                messages.error(request, "Please correct the error below.")
            else:
                messages.error(request, f"Please correct the {error_count} errors below.")

    else:
        # GET request - initialize form with current assignments
        form = FieldJobForm(instance=field_job)
        # Set initial staff members
        current_staff_ids = field_job.assignments.values_list('staff', flat=True)
        form.fields['staff_members'].initial = current_staff_ids

    context = {
        'form': form,
        'field_job': field_job,
        'title': 'Edit Field Job',
        'staff_count': User.objects.filter(role='staff', is_active=True, is_approved=True).count()
    }

    return render(request, 'fieldwork/field_job_form.html', context)

@login_required
@user_passes_test(is_admin)
def field_job_delete(request, pk):
    field_job = get_object_or_404(FieldJob, pk=pk)
    if request.method == 'POST':
        field_job.delete()
        return redirect('fieldwork:field_job_list')
    return render(request, 'fieldwork/field_job_confirm_delete.html', {'field_job': field_job})


@login_required
@user_passes_test(is_admin)
def search_staff(request):
    query = request.GET.get('q', '')
    if query:
        staff_members = User.objects.filter(
            role='staff',
            is_active=True,
            is_approved=True
        ).filter(
            Q(first_name__icontains=query) |
            Q(last_name__icontains=query) |
            Q(email__icontains=query)
        )[:10]

        results = []
        for staff in staff_members:
            results.append({
                'id': staff.id,
                'text': f"{staff.first_name} {staff.last_name} ({staff.email})"
            })

        return JsonResponse({'results': results})
    return JsonResponse({'results': []})


@login_required
@user_passes_test(is_staff)
def staff_field_jobs(request):
    assignments = FieldJobAssignment.objects.filter(staff=request.user).order_by('-created_at')
    return render(request, 'fieldwork/staff_field_jobs.html', {'assignments': assignments})


@login_required
@user_passes_test(is_staff)
def field_report_create(request, assignment_id):
    """
    Field Report creation view. Access is only allowed if check-out is complete.
    """
    print(f"=== DEBUG: field_report_create called ===")
    print(f"Method: {request.method}")
    print(f"Assignment ID: {assignment_id}")
    print(f"User: {request.user}")
    print(f"POST data: {dict(request.POST)}")
    print(f"Navigation check: {request.POST.get('navigation_check', 'NOT FOUND')}")
    print(f"=== END DEBUG ===")

    assignment = get_object_or_404(FieldJobAssignment, id=assignment_id, staff=request.user)

    # Get the existing field report
    try:
        field_report = FieldReport.objects.get(assignment=assignment)
        report_exists = True
    except FieldReport.DoesNotExist:
        field_report = None
        report_exists = False

    # CRITICAL: Check if check-out is complete before allowing access to the form
    if not (field_report and field_report.check_out_time):
        messages.error(request, 'You must check in and check out before creating a field report.')
        return redirect('fieldwork:staff_field_jobs')

    # CRITICAL: If report is already submitted, redirect immediately
    if field_report and field_report.submitted_at:
        messages.info(request, 'Report was already submitted.')
        return redirect('fieldwork:field_report_detail', field_report.id)

    # Handle POST request
    if request.method == 'POST':
        # CRITICAL FIX: Enhanced navigation detection
        navigation_check = request.POST.get('navigation_check', 'false')
        poc_signature = request.POST.get('poc_signature', '').strip()

        # If this is a navigation resubmission, redirect back to form
        if navigation_check == 'true' or (not poc_signature and 'submission_lat' in request.POST):
            logger.warning(
                f"Suspicious form POST (likely intercepted navigation) detected for assignment {assignment_id}. Redirecting back.")
            messages.warning(request,
                             'Form submission was interrupted by navigation. Please use the form controls to submit your report.')
            return redirect('fieldwork:field_report_create', assignment_id=assignment_id)

        # Process the form
        if field_report:
            form = FieldReportForm(request.POST, instance=field_report)
        else:
            form = FieldReportForm(request.POST)

        if form.is_valid():
            # Get location data
            lat = request.POST.get('submission_lat')
            lng = request.POST.get('submission_lng')
            location_address = request.POST.get('submission_address', '')

            if not lat or not lng:
                messages.error(request, 'Location access is required to submit the form.')
            else:
                try:
                    # Validate POC signature
                    if not poc_signature:
                        messages.error(request, 'POC signature is required.')
                    else:
                        expected_code = str(assignment.poc_approval_code)
                        normalized_entered = poc_signature.replace('-', '').lower().strip()
                        normalized_expected = expected_code.replace('-', '').lower()

                        if normalized_entered != normalized_expected:
                            messages.error(request, 'Invalid POC signature.')
                        else:
                            # Valid submission - save the report
                            field_report = form.save(commit=False)
                            field_report.assignment = assignment
                            field_report.submission_location_lat = lat
                            field_report.submission_location_lng = lng
                            field_report.submission_location_address = location_address
                            field_report.submitted_at = timezone.now()

                            # NO CHECKOUT CHECK NEEDED HERE (it's done before accessing the page)

                            field_report.save()
                            assignment.is_code_used = True
                            assignment.save()

                            messages.success(request, 'Field report submitted successfully!')
                            # CRITICAL: Redirect after POST to prevent resubmission
                            return redirect('fieldwork:field_report_detail', field_report.id)

                except Exception as e:
                    logger.error(f"Error saving field report: {str(e)}")
                    messages.error(request, f'An error occurred: {str(e)}')
        else:
            logger.warning(f"Form validation failed: {form.errors}")
            if not poc_signature:
                messages.error(request, 'POC signature is required.')
    else:
        # GET request
        if field_report:
            form = FieldReportForm(instance=field_report)
        else:
            form = FieldReportForm()

    return render(request, 'fieldwork/field_report_form.html', {
        'assignment': assignment,
        'form': form,
        'report_exists': report_exists,
        'field_report': field_report
    })


# fieldwork/views.py (Add new imports at the top)
# ... existing imports
from io import BytesIO # Keep this for memory buffer
from docx import Document # 🆕 DOCX Library
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.shared import RGBColor
# ... existing imports
from django.http import JsonResponse, HttpResponse # <--- Add HttpResponse here
# ----------------------------------------------------------------------
# DOCX Export Function using python-docx
# ----------------------------------------------------------------------
@login_required
@user_passes_test(lambda u: u.role == 'admin')
def field_report_export_docx(request, report_id):
    """
    Generates the Field Report as a DOCX file, retrieves associated data,
    and forces an automatic download.
    """
    field_report = get_object_or_404(FieldReport, id=report_id)
    job = field_report.assignment.field_job
    staff = field_report.assignment.staff

    document = Document()

    # --- Document Styling Setup ---
    # Set default font size for the entire document
    style = document.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # --- 1. Title and Header ---

    # Title (Large and centered)
    title = document.add_heading(f"Field Report: {job.client_name}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle/Metadata
    document.add_paragraph(f"Report ID: {report_id} | Submitted By: {staff.get_full_name()} (FE)", style='Subtitle')
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(
        f"Submission Date: {field_report.submitted_at.strftime('%Y-%m-%d %H:%M:%S') if field_report.submitted_at else 'N/A'}")
    document.add_paragraph()  # Spacer

    # --- Helper function for adding styled data points ---
    def add_data_point(doc, label, value):
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(str(value) or "N/A")

    # --- 2. Job & Personnel Information ---
    document.add_heading('1. Job & Personnel Information', level=1)

    add_data_point(document, "Client Name", job.client_name)
    add_data_point(document, "Category", job.category)
    add_data_point(document, "Job Date/Time", job.date_of_job.strftime('%Y-%m-%d %H:%M:%S'))
    add_data_point(document, "Site Address", job.site_address)
    add_data_point(document, "POC Name", job.poc_name)
    add_data_point(document, "POC Email", job.poc_email)
    add_data_point(document, "Field Engineer", staff.get_full_name())
    document.add_paragraph()

    # --- 3. Time Tracking ---
    document.add_heading('2. Time Tracking', level=1)

    add_data_point(document, "Check In Time",
                   field_report.check_in_time.strftime('%Y-%m-%d %H:%M:%S') if field_report.check_in_time else 'N/A')
    add_data_point(document, "Check Out Time",
                   field_report.check_out_time.strftime('%Y-%m-%d %H:%M:%S') if field_report.check_out_time else 'N/A')
    add_data_point(document, "Total Hours Worked",
                   f"{field_report.hours_worked} hrs" if field_report.hours_worked else 'N/C')
    document.add_paragraph()

    # --- 4. Work Summary ---
    document.add_heading('3. Work Summary', level=1)

    # Original Scope
    document.add_paragraph().add_run("Original Scope of Work:").bold = True
    document.add_paragraph(job.summary_of_work or "Not specified.", style='List Bullet')
    document.add_paragraph()

    # Work Performed
    document.add_paragraph().add_run("Work Performed by Field Engineer:").bold = True
    document.add_paragraph(field_report.summary_of_work_performed or "No summary provided.", style='List Bullet')
    document.add_paragraph()

    # --- 5. Final Status & Signatures ---
    document.add_heading('4. Final Status & Signatures', level=1)

    # Status Points (with conditional formatting/text)
    add_data_point(document, "Site Visit Complete", "Yes" if field_report.site_visit_complete else "No")
    add_data_point(document, "Revisit Required", "Yes" if field_report.revisit_required else "No")
    add_data_point(document, "POC Code Used", "Yes" if field_report.assignment.is_code_used else "No")

    # FE Signature
    sig_p = document.add_paragraph()
    sig_p.add_run("FE Signature: ").bold = True
    sig_p.add_run(field_report.fe_signature or "Unsigned")
    document.add_paragraph()

    # Location
    document.add_heading('5. Submission Location', level=1)
    add_data_point(document, "Submission Address", field_report.submission_location_address)
    add_data_point(document, "Coordinates",
                   f"Lat: {field_report.submission_location_lat}, Lng: {field_report.submission_location_lng}")

    # --- 6. Save document to a buffer and trigger download ---

    # Use BytesIO to save the document to memory
    doc_io = BytesIO()
    document.save(doc_io)
    doc_io.seek(0)

    # Filename
    file_name = f"Field_Report_{job.client_name.replace(' ', '_')}_{report_id}.docx"

    # Create the HTTP response for download
    response = HttpResponse(
        doc_io.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    # Set Content-Disposition to 'attachment' to force a download prompt
    response['Content-Disposition'] = f'attachment; filename="{file_name}"'

    return response

@login_required
def check_in(request, assignment_id):
    if request.method == 'POST':
        assignment = get_object_or_404(FieldJobAssignment, id=assignment_id, staff=request.user)

        # Get or create field report
        field_report, created = FieldReport.objects.get_or_create(assignment=assignment)

        if not field_report.check_in_time:
            field_report.check_in_time = timezone.now()
            field_report.save()
            return JsonResponse({'success': True, 'check_in_time': field_report.check_in_time.isoformat()})
        else:
            return JsonResponse({'success': False, 'error': 'Already checked in'})

    return JsonResponse({'success': False, 'error': 'Invalid request'})


@login_required
def check_out(request, assignment_id):
    if request.method == 'POST':
        assignment = get_object_or_404(FieldJobAssignment, id=assignment_id, staff=request.user)

        try:
            field_report = FieldReport.objects.get(assignment=assignment)

            if not field_report.check_in_time:
                return JsonResponse({'success': False, 'error': 'Please check in first'})

            if not field_report.check_out_time:
                field_report.check_out_time = timezone.now()

                # Calculate hours worked
                time_diff = field_report.check_out_time - field_report.check_in_time
                hours_worked = time_diff.total_seconds() / 3600
                field_report.hours_worked = round(hours_worked, 2)

                field_report.save()
                return JsonResponse({'success': True, 'check_out_time': field_report.check_out_time.isoformat()})
            else:
                return JsonResponse({'success': False, 'error': 'Already checked out'})
        except FieldReport.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Please check in first'})

    return JsonResponse({'success': False, 'error': 'Invalid request'})

@login_required
def field_report_detail(request, report_id):
    field_report = get_object_or_404(FieldReport, id=report_id)

    # Check permissions
    if not (request.user.role == 'admin' or field_report.assignment.staff == request.user):
        return redirect('fieldwork:staff_field_jobs')

    return render(request, 'fieldwork/field_report_detail.html', {'field_report': field_report})


def send_poc_approval_email(field_job, assignment):
    """Send approval code email to POC for a specific assignment"""
    try:
        subject = f'Field Job Approval Code - {field_job.client_name}'

        # Safe way to get staff name with multiple fallbacks
        try:
            staff_name = assignment.staff.get_full_name()
        except AttributeError:
            # Fallback to manual name construction
            staff_name = f"{assignment.staff.first_name} {assignment.staff.last_name}".strip()
            if not staff_name:
                staff_name = assignment.staff.email.split('@')[0]  # Use email username as fallback
            if not staff_name:
                staff_name = "Field Engineer"  # Final fallback

        # Format the datetime for display in email
        job_datetime = timezone.localtime(field_job.date_of_job)
        formatted_datetime = job_datetime.strftime('%B %d, %Y at %I:%M %p')

        context = {
            'field_job': field_job,
            'assignment': assignment,
            'staff_name': staff_name,
            'formatted_datetime': formatted_datetime,
        }

        message = render_to_string('fieldwork/emails/poc_approval_email.html', context)

        send_mail(
            subject,
            f'Please use approval code {assignment.poc_approval_code} for {staff_name} at {field_job.client_name} on {formatted_datetime}.',
            settings.DEFAULT_FROM_EMAIL,
            [field_job.poc_email],
            html_message=message,
            fail_silently=False,
        )

        logger.info(f"POC approval email sent to {field_job.poc_email} for assignment {assignment.id}")

    except Exception as e:
        logger.error(f"Failed to send POC approval email: {str(e)}")
        raise


def send_staff_assignment_email(field_job, staff):
    """Send assignment notification email to staff member"""
    try:
        subject = f'New Field Job Assignment - {field_job.client_name}'

        # Safe way to get staff name
        try:
            staff_name = staff.get_full_name()
        except AttributeError:
            staff_name = f"{staff.first_name} {staff.last_name}".strip() or staff.email

        # Format the datetime for display in email
        job_datetime = timezone.localtime(field_job.date_of_job)
        formatted_datetime = job_datetime.strftime('%B %d, %Y at %I:%M %p')

        context = {
            'field_job': field_job,
            'staff': staff,
            'staff_name': staff_name,
            'formatted_datetime': formatted_datetime,
            'protocol': 'https' if settings.DEBUG else 'https',
            'domain': settings.ALLOWED_HOSTS[0] if settings.ALLOWED_HOSTS else 'localhost:8000',
        }

        message = render_to_string('fieldwork/emails/staff_assignment_email.html', context)

        send_mail(
            subject,
            f'You have been assigned to a field job at {field_job.client_name} on {formatted_datetime}.',
            settings.DEFAULT_FROM_EMAIL,
            [staff.email],
            html_message=message,
            fail_silently=False,
        )

        logger.info(f"Staff assignment email sent to {staff.email} for job {field_job.id}")

    except Exception as e:
        logger.error(f"Failed to send staff assignment email: {str(e)}")
        raise


@login_required
@require_GET
def job_details_api(request, job_id):
    """API endpoint to get job details including approval codes"""
    try:
        field_job = FieldJob.objects.get(id=job_id)

        # Format the datetime for display
        job_datetime = timezone.localtime(field_job.date_of_job)
        formatted_datetime = job_datetime.strftime('%B %d, %Y at %I:%M %p')

        # Prepare assignments data
        assignments_data = []
        for assignment in field_job.assignments.all():
            assignments_data.append({
                'staff_name': assignment.staff.get_full_name(),
                'approval_code': str(assignment.poc_approval_code)  # This is the REAL code
            })

        data = {
            'id': field_job.id,
            'client_name': field_job.client_name,
            'category': field_job.category,
            'date_of_job': field_job.date_of_job.isoformat(),
            'formatted_datetime': formatted_datetime,
            'site_address': field_job.site_address,
            'poc_name': field_job.poc_name,
            'summary_of_work': field_job.summary_of_work,
            'assignments': assignments_data
        }

        return JsonResponse(data)

    except FieldJob.DoesNotExist:
        return JsonResponse({'error': 'Job not found'}, status=404)


def send_staff_update_email(field_job, staff):
    """Send update notification email to staff member"""
    try:
        subject = f'Field Job Updated - {field_job.client_name}'

        # Safe way to get staff name
        try:
            staff_name = staff.get_full_name()
        except AttributeError:
            staff_name = f"{staff.first_name} {staff.last_name}".strip() or staff.email

        # Format the datetime for display in email
        job_datetime = timezone.localtime(field_job.date_of_job)
        formatted_datetime = job_datetime.strftime('%B %d, %Y at %I:%M %p')

        context = {
            'field_job': field_job,
            'staff': staff,
            'staff_name': staff_name,
            'formatted_datetime': formatted_datetime,
            'protocol': 'https' if settings.DEBUG else 'https',
            'domain': settings.ALLOWED_HOSTS[0] if settings.ALLOWED_HOSTS else 'localhost:8000',
            'is_update': True,
        }

        message = render_to_string('fieldwork/emails/staff_update_email.html', context)

        send_mail(
            subject,
            f'Your field job assignment at {field_job.client_name} on {formatted_datetime} has been updated.',
            settings.DEFAULT_FROM_EMAIL,
            [staff.email],
            html_message=message,
            fail_silently=False,
        )

        logger.info(f"Staff update email sent to {staff.email} for job {field_job.id}")

    except Exception as e:
        logger.error(f"Failed to send staff update email: {str(e)}")
        raise

def send_poc_update_email(field_job, assignment):
    """Send update notification email to POC for a specific assignment"""
    try:
        subject = f'Field Job Updated - {field_job.client_name}'

        # Safe way to get staff name with multiple fallbacks
        try:
            staff_name = assignment.staff.get_full_name()
        except AttributeError:
            # Fallback to manual name construction
            staff_name = f"{assignment.staff.first_name} {assignment.staff.last_name}".strip()
            if not staff_name:
                staff_name = assignment.staff.email.split('@')[0]  # Use email username as fallback
            if not staff_name:
                staff_name = "Field Engineer"  # Final fallback

        # Format the datetime for display in email
        job_datetime = timezone.localtime(field_job.date_of_job)
        formatted_datetime = job_datetime.strftime('%B %d, %Y at %I:%M %p')

        context = {
            'field_job': field_job,
            'assignment': assignment,
            'staff_name': staff_name,
            'formatted_datetime': formatted_datetime,
            'is_update': True,
        }

        message = render_to_string('fieldwork/emails/poc_update_email.html', context)

        send_mail(
            subject,
            f'Field job for {field_job.client_name} with {staff_name} on {formatted_datetime} has been updated.',
            settings.DEFAULT_FROM_EMAIL,
            [field_job.poc_email],
            html_message=message,
            fail_silently=False,
        )

        logger.info(f"POC update email sent to {field_job.poc_email} for assignment {assignment.id}")

    except Exception as e:
        logger.error(f"Failed to send POC update email: {str(e)}")
        raise


@login_required
@user_passes_test(is_admin)
def field_job_evaluation(request, pk):
    """
    Evaluation page for a field job - shows all submitted reports from assigned staff
    """
    field_job = get_object_or_404(FieldJob, pk=pk)

    if not request.user.role == 'admin':
        raise PermissionDenied("You don't have permission to evaluate field jobs.")

    # Get all assignments for this job with their reports
    assignments_with_reports = FieldJobAssignment.objects.filter(
        field_job=field_job
    ).select_related('field_report').order_by('staff__first_name')

    # Handle search functionality
    search_query = request.GET.get('search', '')
    if search_query:
        assignments_with_reports = assignments_with_reports.filter(
            Q(staff__first_name__icontains=search_query) |
            Q(staff__last_name__icontains=search_query) |
            Q(field_report__summary_of_work_performed__icontains=search_query)
        )

    context = {
        'field_job': field_job,
        'assignments_with_reports': assignments_with_reports,
        'search_query': search_query,
    }

    return render(request, 'fieldwork/field_job_evaluation.html', context)


import requests
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_GET
import logging

logger = logging.getLogger(__name__)


@login_required
@require_GET
def reverse_geocode(request):
    """
    Reverse geocoding endpoint that converts latitude/longitude to address
    using OpenStreetMap Nominatim (FREE)
    """
    # Get latitude and longitude from request parameters
    lat = request.GET.get('lat')
    lng = request.GET.get('lng')

    # Validate parameters
    if not lat or not lng:
        return JsonResponse({
            'success': False,
            'error': 'Latitude and longitude are required parameters'
        }, status=400)

    try:
        # Validate coordinates are numbers
        lat_float = float(lat)
        lng_float = float(lng)

        # Validate coordinate ranges
        if not (-90 <= lat_float <= 90) or not (-180 <= lng_float <= 180):
            return JsonResponse({
                'success': False,
                'error': 'Invalid coordinate values'
            }, status=400)

    except ValueError:
        return JsonResponse({
            'success': False,
            'error': 'Invalid coordinate format'
        }, status=400)

    try:
        # OpenStreetMap Nominatim API endpoint
        url = f"https://nominatim.openstreetmap.org/reverse"

        # Parameters for the request
        params = {
            'format': 'json',
            'lat': lat,
            'lon': lng,
            'zoom': 18,  # Level of detail (18 = building level)
            'addressdetails': 1  # Get detailed address components
        }

        # Headers as required by Nominatim usage policy
        headers = {
            'User-Agent': 'YourCompanyFieldApp/1.0 (contact@yourcompany.com)',  # CHANGE THIS!
            'Accept-Language': 'en'
        }
        # Make the request to OpenStreetMap
        logger.info(f"Making reverse geocode request for coordinates: {lat}, {lng}")
        response = requests.get(url, params=params, headers=headers, timeout=10)

        # Check if request was successful
        if response.status_code == 200:
            data = response.json()

            # Extract address information
            if data and 'display_name' in data:
                address = data['display_name']

                # You can also extract specific address components if needed
                address_components = {}
                if 'address' in data:
                    address_data = data['address']
                    address_components = {
                        'road': address_data.get('road', ''),
                        'suburb': address_data.get('suburb', ''),
                        'city': address_data.get('city', ''),
                        'state': address_data.get('state', ''),
                        'postcode': address_data.get('postcode', ''),
                        'country': address_data.get('country', ''),
                    }

                logger.info(f"Successfully geocoded coordinates to: {address}")

                return JsonResponse({
                    'success': True,
                    'address': address,
                    'components': address_components,
                    'raw_data': data  # Include full response for debugging
                })
            else:
                logger.warning(f"No address found for coordinates: {lat}, {lng}")
                return JsonResponse({
                    'success': True,
                    'address': f"Location at {lat:.6f}, {lng:.6f}",
                    'components': {},
                    'note': 'No specific address found, using coordinates'
                })
        else:
            logger.error(f"OpenStreetMap API error: {response.status_code}")
            return JsonResponse({
                'success': False,
                'error': f'Geocoding service error: {response.status_code}'
            }, status=500)

    except requests.exceptions.Timeout:
        logger.error("Geocoding request timed out")
        return JsonResponse({
            'success': False,
            'error': 'Geocoding service timeout'
        }, status=504)

    except requests.exceptions.RequestException as e:
        logger.error(f"Geocoding request failed: {str(e)}")
        return JsonResponse({
            'success': False,
            'error': 'Geocoding service unavailable'
        }, status=503)

    except Exception as e:
        logger.error(f"Unexpected error in reverse_geocode: {str(e)}")
        return JsonResponse({
            'success': False,
            'error': 'Internal server error'
        }, status=500)


@login_required
@user_passes_test(is_admin)
def get_report_data(request, report_id):
    """API endpoint to get report data for editing"""
    field_report = get_object_or_404(FieldReport, id=report_id)

    return JsonResponse({
        'success': True,
        'report': {
            'summary_of_work_performed': field_report.summary_of_work_performed,
            'site_visit_complete': field_report.site_visit_complete,
            'revisit_required': field_report.revisit_required,
            'hours_worked': field_report.hours_worked,
            'fe_signature': field_report.fe_signature,
            'submission_location_address': field_report.submission_location_address,
            'submission_location_lat': field_report.submission_location_lat,
            'submission_location_lng': field_report.submission_location_lng,
        }
    })


@login_required
@user_passes_test(is_admin)
def update_report(request, report_id):
    """API endpoint to update report data"""
    if request.method == 'POST':
        field_report = get_object_or_404(FieldReport, id=report_id)

        # Update editable fields (excluding location data)
        field_report.summary_of_work_performed = request.POST.get('summary', '')
        field_report.site_visit_complete = request.POST.get('site_visit_complete') == 'on'
        field_report.revisit_required = request.POST.get('revisit_required') == 'on'
        field_report.hours_worked = request.POST.get('hours_worked')
        field_report.fe_signature = request.POST.get('fe_signature', '')

        field_report.save()

        return JsonResponse({'success': True})

    return JsonResponse({'success': False, 'error': 'Invalid request method'})


@login_required
@user_passes_test(is_admin)
def delete_report(request, report_id):
    """API endpoint to delete a report"""
    if request.method == 'POST':
        field_report = get_object_or_404(FieldReport, id=report_id)
        field_report.delete()

        return JsonResponse({'success': True})

    return JsonResponse({'success': False, 'error': 'Invalid request method'})


from .services.ipapi_service import IPAPIService
import json


# Add this new view for IPAPI location
@login_required
@require_GET
def get_ipapi_location(request):
    """
    Primary location endpoint using IPAPI.co
    Falls back to browser geolocation if IPAPI fails
    """
    try:
        # Try IPAPI first
        ipapi_location = IPAPIService.get_client_location(request)

        if ipapi_location:
            if ipapi_location.get('error') == 'quota_exceeded':
                logger.warning("IPAPI quota exceeded, will use fallback")
                return JsonResponse({
                    'success': False,
                    'error': 'ipapi_quota_exceeded',
                    'message': 'IP location service quota exceeded, using browser location'
                })

            return JsonResponse({
                'success': True,
                'source': 'ipapi',
                'location': {
                    'lat': ipapi_location['lat'],
                    'lng': ipapi_location['lng'],
                    'address': ipapi_location['address'],
                    'accuracy': ipapi_location['accuracy'],
                    'city': ipapi_location.get('city'),
                    'country': ipapi_location.get('country')
                }
            })
        else:
            # IPAPI failed, indicate to use browser geolocation
            return JsonResponse({
                'success': False,
                'error': 'ipapi_failed',
                'message': 'IP location unavailable, using browser geolocation'
            })

    except Exception as e:
        logger.error(f"Error in IPAPI location endpoint: {str(e)}")
        return JsonResponse({
            'success': False,
            'error': 'server_error',
            'message': 'Location service error'
        })


# Update your existing reverse_geocode view to be fallback
@login_required
@require_GET
def reverse_geocode(request):
    """
    Fallback reverse geocoding (now secondary to IPAPI)
    """
    lat = request.GET.get('lat')
    lng = request.GET.get('lng')

    if not lat or not lng:
        return JsonResponse({
            'success': False,
            'error': 'Latitude and longitude are required'
        }, status=400)

    try:
        # Your existing OpenStreetMap code here
        url = "https://nominatim.openstreetmap.org/reverse"
        params = {
            'format': 'json',
            'lat': lat,
            'lon': lng,
            'zoom': 18,
            'addressdetails': 1
        }
        headers = {
            'User-Agent': 'FieldWorkApp/1.0 (your-email@domain.com)',
            'Accept-Language': 'en'
        }

        response = requests.get(url, params=params, headers=headers, timeout=10)

        if response.status_code == 200:
            data = response.json()
            if data and 'display_name' in data:
                address = data['display_name']
                return JsonResponse({
                    'success': True,
                    'source': 'fallback_geocoding',
                    'address': address
                })

        return JsonResponse({
            'success': False,
            'error': 'Geocoding failed'
        })

    except Exception as e:
        logger.error(f"Fallback geocoding failed: {str(e)}")
        return JsonResponse({
            'success': False,
            'error': 'Geocoding service unavailable'
        })



