import base64
from .models import RFPReferral, RFPFile, OriginalDocxFile  # Add OriginalDocxFile
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib import messages
from django.http import HttpResponseForbidden, HttpResponseRedirect, JsonResponse
from django.db.models import Q
from django.core.mail import send_mail
from django.template.loader import render_to_string
from django.utils import timezone
import os
import cloudinary.uploader
import re  # Added for postal code validation

from .models import RFPReferral, RFPFile
from django.contrib.auth import get_user_model

User = get_user_model()


# --- Decorator for Admin Permission ---
def admin_required(function=None, redirect_field_name=None, login_url='/accounts/login/'):
    """Checks if the user is authenticated AND has role 'admin'."""
    actual_decorator = user_passes_test(
        lambda u: u.is_active and hasattr(u, 'role') and u.role == 'admin',
        login_url=login_url,
        redirect_field_name=redirect_field_name
    )
    if function:
        return actual_decorator(function)
    return actual_decorator


def send_rfp_submission_email(referral):
    """Sends a notification email to the administrative contact upon initial submission."""

    TO_EMAIL = ['hyallison5050@gmail.com']  # Target admin email
    FROM_EMAIL = 'petermomo124@gmail.com'  # Sender email

    subject = f'New RFP Submission: {referral.proposal_title}'

    # Ensure you have rfp_submission_email.html in rfp/templates/rfp/
    html_message = render_to_string('rfp/rfp_submission_email.html', {
        'referral': referral,
    })

    plain_message = f"""
    A new Request for Proposal has been submitted.

    Title: {referral.proposal_title}
    Organization: {referral.organization}
    Deadline: {referral.proposal_deadline}
    Client Email: {referral.client.email}
    """

    try:
        send_mail(
            subject,
            plain_message,
            FROM_EMAIL,
            TO_EMAIL,
            html_message=html_message,
            fail_silently=False,
        )
    except Exception as e:
        print(f"Failed to send RFP submission email: {e}")


def send_rfp_update_notification(referral, user_who_made_change, action):
    """
    Sends a dynamic notification email to the counter-party based on the 'action' taken.

    Actions: 'update', 'file_add', 'file_delete', 'delete'.
    NOTE: This uses the 'emails/rfp_update_notification.html' template.

    MODIFIED: Returns True on success, False on failure.
    """

    ADMIN_EMAIL = ['hyallison5050@gmail.com']
    FROM_EMAIL = 'petermomo124@gmail.com'

    is_admin = hasattr(user_who_made_change, 'role') and user_who_made_change.role == 'admin'

    # --- Define Subject and Instructions based on Action and User Role ---
    if not is_admin:  # Client made the change (Notify Admin)
        recipient_list = ADMIN_EMAIL
        notification_type = "Client Action"

        if action == 'update':
            subject = f"ACTION REQUIRED: Client Updated RFP Details - {referral.proposal_title}"
            instructions = "A client has edited the main details of the RFP below. Please review the changes in the portal."
        elif action == 'file_add':
            subject = f"FILE ADDED: Client Uploaded File to RFP - {referral.proposal_title}"
            instructions = "A client has uploaded new documents to the RFP. Please log in to download and review the attachments."
        elif action == 'file_delete':
            subject = f"FILE DELETED: Client Removed File from RFP - {referral.proposal_title}"
            instructions = "A client has removed one of their uploaded files from the RFP. Please confirm."
        else:  # Should not happen, fallback
            subject = f"UNSPECIFIED CLIENT ACTION on RFP - {referral.proposal_title}"
            instructions = "A client has performed an unspecific action on the RFP. Please investigate."

    else:  # Admin made the change (Notify Client)
        recipient_list = [referral.client.email]
        notification_type = "Admin Action"

        if action == 'update':
            subject = f"STATUS/DETAIL UPDATE: Your RFP Has Been Modified - {referral.proposal_title}"
            instructions = f"An administrator has made an update to your RFP (Status: {referral.status}). Please log into your client portal for details."
        elif action == 'file_add':
            subject = f"FILE ADDED: Admin Uploaded File to RFP - {referral.proposal_title}"
            instructions = "An administrator has uploaded new documents related to your RFP. Please log in to view the new attachments."
        elif action == 'file_delete':
            subject = f"FILE DELETED: Admin Removed File from RFP - {referral.proposal_title}"
            instructions = "An administrator has removed a file from your RFP documents. Please log in to verify."
        elif action == 'delete':
            # This is sent right before deletion
            subject = f"RFP DELETED: Your Request for Proposal Has Been Removed"
            instructions = "An administrator has removed the following RFP from the system. If you believe this is an error, please contact support."

        # Add this condition to your existing function
        elif action == 'file_edit':
            if not is_admin:
                subject = f"DOCUMENT EDITED: Client Modified File - {referral.proposal_title}"
                instructions = "A client has edited a document in the RFP. Please review the changes."
            else:
                subject = f"DOCUMENT EDITED: Admin Modified File - {referral.proposal_title}"
                instructions = "An administrator has edited a document in your RFP. Please review the changes."
        else:  # Should not happen, fallback
            subject = f"UNSPECIFIED ADMIN ACTION on RFP - {referral.proposal_title}"
            instructions = "An administrator has performed an unspecific action on your RFP. Please log in for details."

    # --- Render and Send Email ---
    html_content = render_to_string('emails/rfp_update_notification.html', {
        'referral': referral,
        'notification_type': notification_type,
        'instructions': instructions,
        'recipient_is_admin': not is_admin,  # If True, the recipient is the client.
        'user_who_made_change': user_who_made_change.email,
        'action': action,  # <--- ADDED: Pass the action type to the template
    })

    plain_message = f"""
    RFP {referral.proposal_title} has been updated.
    Action: {notification_type}
    Last Edited By: {user_who_made_change.email}
    """

    try:
        send_mail(
            subject=subject,
            message=plain_message,
            from_email=FROM_EMAIL,
            recipient_list=recipient_list,
            html_message=html_content,
            fail_silently=False,
        )
        return True  # Return success status
    except Exception as e:
        print(f"Error sending RFP update notification email: {e}")
        return False  # Return failure status




# ====================================================================
# VIEWS: CLIENT AND ADMIN
# ====================================================================

@login_required
def rfp_list(request):
    """Lists RFP referrals based on user role."""
    user = request.user

    is_admin = hasattr(user, 'role') and user.role == 'admin'

    # Admin sees all referrals
    if is_admin:
        referrals = RFPReferral.objects.all().order_by('-created_at')
    # Client sees only their own
    else:
        # This handles 'client' role and any other non-admin role
        referrals = RFPReferral.objects.filter(client=user).order_by('-created_at')

    context = {
        'referrals': referrals,
        'is_admin': is_admin,
    }
    return render(request, 'rfp/rfp_list.html', context)


@login_required
def rfp_create(request):
    """
    Handles creation of a new RFP referral with submission notification.
    The submission email is already designed to be the first notification.
    """
    user = request.user

    # Client role check
    if hasattr(user, 'role') and user.role != 'client':
        messages.error(request, "Only clients are permitted to create new RFP referrals.")
        return redirect('rfp_list')

    if request.method == 'POST':
        try:
            # 1. Retrieve ALL required fields from POST data
            proposal_title = request.POST.get('proposal_title')
            organization = request.POST.get('organization')
            deadline_str = request.POST.get('proposal_deadline')
            city = request.POST.get('city')
            state_province = request.POST.get('state_province')
            postal_code = request.POST.get('postal_code')

            # 2. Server-side validation for empty fields
            required_fields = {
                'Proposal Title': proposal_title,
                'Organization': organization,
                'Deadline': deadline_str,
                'City': city,
                'State/Province': state_province,
                'Postal/Zip Code': postal_code,
            }

            missing_fields = [name for name, value in required_fields.items() if not value]
            if missing_fields:
                messages.error(request, f'The following fields are required: {", ".join(missing_fields)}.')
                return redirect('rfp_create')

            # 3. Custom Validation: Postal Code Regex Check
            postal_code_pattern = re.compile(r'^[a-zA-Z0-9\s()+-]*$')

            if not postal_code_pattern.match(postal_code):
                messages.error(request,
                               'Invalid Postal/Zip Code format. Only letters, numbers, spaces, and the symbols ( ) - + are allowed.')
                return redirect('rfp_create')

            # 4. Date validation and conversion
            deadline = timezone.datetime.strptime(deadline_str, '%Y-%m-%dT%H:%M')

            # 5. Create the referral
            referral = RFPReferral.objects.create(
                proposal_title=proposal_title,
                organization=organization,
                proposal_deadline=deadline,
                city=city,
                state_province=state_province,
                postal_code=postal_code,
                client=request.user,
                status='submitted'
            )

            # Send Email Notification (RFP Submission)
            send_rfp_submission_email(referral)

            messages.success(request,
                             'RFP referral submitted successfully! An email notification has been sent. You can now attach files.')
            return redirect('rfp_detail', pk=referral.pk)

        except ValueError:
            messages.error(request, 'Invalid date/time format for the proposal deadline.')
        except Exception as e:
            messages.error(request, f'An unexpected error occurred during submission: {e}')

        return redirect('rfp_create')

    # If GET request
    return render(request, 'rfp/rfp_create.html')


@login_required
def rfp_detail(request, pk):
    """
    Displays RFP detail and file upload section, filtering files into
    client-uploaded and admin-uploaded sections.
    """
    referral = get_object_or_404(RFPReferral, pk=pk)
    user = request.user

    is_admin = hasattr(user, 'role') and user.role == 'admin'

    # Permission check: must be admin OR the client who submitted it
    if not (is_admin or referral.client == user):
        return HttpResponseForbidden("You do not have permission to view this referral.")

    # Fetch all files related to the referral
    all_files = referral.files.all()

    # 1. Filter files uploaded by the original client
    client_files = all_files.filter(uploaded_by=referral.client).order_by('-id')

    # 2. Filter files uploaded by any user with the 'admin' role
    admin_files = all_files.filter(uploaded_by__role='admin').order_by('-id')

    context = {
        'referral': referral,
        'is_admin': is_admin,
        'client_files': client_files,
        'admin_files': admin_files,
        'allowed_extensions_list': ['.doc', '.docx']
    }
    return render(request, 'rfp/rfp_detail.html', context)


@login_required
def rfp_update(request, pk):
    """
    Allows client/admin to edit their referral, triggering a notification to the counter-party.
    """
    referral = get_object_or_404(RFPReferral, pk=pk)
    user = request.user

    is_admin = hasattr(user, 'role') and user.role == 'admin'

    # Permission check: must be admin OR the client who submitted it
    if not (is_admin or referral.client == user):
        return HttpResponseForbidden("You do not have permission to update this referral.")

    if request.method == 'POST':
        try:
            # We don't track old status here, notification is sent regardless of content change

            referral.proposal_title = request.POST.get('proposal_title')
            referral.organization = request.POST.get('organization')

            # Date validation and conversion
            deadline_str = request.POST.get('proposal_deadline')
            referral.proposal_deadline = timezone.datetime.strptime(deadline_str, '%Y-%m-%dT%H:%M')

            referral.city = request.POST.get('city')
            referral.state_province = request.POST.get('state_province')
            referral.postal_code = request.POST.get('postal_code')

            # Admin can edit status
            if is_admin:
                referral.status = request.POST.get('status')

            referral.save()

            # --- INTEGRATION: Send Update Notification ---
            send_rfp_update_notification(referral, user, action='update')
            # --------------------------------------------

            messages.success(request, 'RFP referral updated successfully. Notification sent.')
            return redirect('rfp_detail', pk=referral.pk)

        except ValueError:
            messages.error(request, 'Invalid date/time format for the proposal deadline.')
        except Exception as e:
            messages.error(request, f'An unexpected error occurred during update: {e}')

        return redirect('rfp_update', pk=referral.pk)

    context = {
        'referral': referral,
        'is_admin': is_admin,
        'status_choices': RFPReferral.STATUS_CHOICES,
    }
    return render(request, 'rfp/rfp_update.html', context)


@admin_required
def rfp_delete(request, pk):
    """
    Admin-only view to delete a referral, notifying the client.
    MODIFIED: Captures and reports email sending status.
    """
    referral = get_object_or_404(RFPReferral, pk=pk)

    if request.method == 'POST':

        # --- INTEGRATION: Send Delete Notification BEFORE deleting the object ---
        email_sent = send_rfp_update_notification(referral, request.user, action='delete')
        # -----------------------------------------------------------------------

        referral.delete()

        if email_sent:
            messages.success(request, 'RFP referral successfully deleted. Client notified.')
        else:
            messages.warning(request,
                             'RFP referral successfully deleted, but the client notification email failed to send.')

        return redirect('rfp_list')

    return render(request, 'rfp/rfp_delete.html', {'referral': referral})


# ====================================================================
# VIEWS: FILE MANAGEMENT
# ====================================================================
@login_required
def rfp_add_file(request, rfp_pk):
    """
    Handles file upload and saves to Cloudinary, notifying the counter-party.
    """
    referral = get_object_or_404(RFPReferral, pk=rfp_pk)
    user = request.user

    is_admin = hasattr(user, 'role') and user.role == 'admin'
    if not (is_admin or referral.client == user):
        return HttpResponseForbidden("You do not have permission to add files to this referral.")

    ALLOWED_EXTENSIONS = ['.doc', '.docx']

    if request.method == 'POST':
        uploaded_files = request.FILES.getlist('files_to_upload')

        if not uploaded_files:
            messages.error(request, 'Please select one or more files to upload.')
            return redirect('rfp_detail', pk=referral.pk)

        successful_uploads = 0

        for uploaded_file in uploaded_files:
            file_name = uploaded_file.name
            _, file_extension = os.path.splitext(file_name)
            file_extension_lower = file_extension.lower()

            if file_extension_lower not in ALLOWED_EXTENSIONS:
                messages.error(
                    request,
                    f"File '{file_name}' type ('{file_extension}') is not allowed. Only {', '.join(ALLOWED_EXTENSIONS)} are permitted."
                )
                continue

            upload_resource_type = 'raw'

            try:
                folder_path = f"rfp/{rfp_pk}/attachments"

                upload_result = cloudinary.uploader.upload(
                    uploaded_file,
                    folder=folder_path,
                    resource_type=upload_resource_type,
                    access_mode='public',
                    timeout=600
                )

                RFPFile.objects.create(
                    referral=referral,
                    cloudinary_url=upload_result.get('secure_url'),
                    public_id=upload_result.get('public_id'),
                    uploaded_by=user,
                )
                successful_uploads += 1

            except Exception as e:
                messages.error(request,
                               f"Upload failed for '{file_name}': A server error occurred. ({type(e).__name__})")

        if successful_uploads > 0:
            # --- INTEGRATION: Send File Add Notification ---
            email_sent = send_rfp_update_notification(referral, user, action='file_add')
            # -----------------------------------------------

            # MODIFIED: Check email_sent status and display appropriate message
            if email_sent:
                messages.success(request, f"{successful_uploads} file(s) added successfully. Counter-party notified.")
            else:
                messages.warning(request,
                                 f"{successful_uploads} file(s) added successfully, but the notification email to the counter-party failed. Check system logs for details.")

        return redirect('rfp_detail', pk=referral.pk)

    return redirect('rfp_detail', pk=referral.pk)


@login_required
def rfp_delete_file(request, rfp_pk, file_pk):
    """
    Allows Admin to delete any file. Allows Client to delete their own files, notifying the counter-party.
    """
    referral = get_object_or_404(RFPReferral, pk=rfp_pk)
    rfp_file = get_object_or_404(RFPFile, pk=file_pk, referral=referral)
    user = request.user

    is_admin = hasattr(user, 'role') and user.role == 'admin'

    # Permission check: Admin can delete any. Client can only delete their own uploaded files.
    if not (is_admin or (referral.client == user and rfp_file.uploaded_by == user)):
        return HttpResponseForbidden("You do not have permission to delete this file.")

    filename = os.path.basename(rfp_file.public_id).split('/')[-1] if rfp_file.public_id else "File"

    if request.method == 'POST':
        try:
            # 1. Delete from Cloudinary
            cloudinary.uploader.destroy(rfp_file.public_id)

            # 2. Delete from Database
            rfp_file.delete()

            # --- INTEGRATION: Send File Delete Notification ---
            send_rfp_update_notification(referral, user, action='file_delete')
            # --------------------------------------------------

            messages.success(request, f'File "{filename}" deleted successfully. Counter-party notified.')

        except Exception as e:
            # Handle cases where Cloudinary deletion might fail but we proceed with DB deletion
            rfp_file.delete()
            messages.warning(request,
                             f'File "{filename}" was deleted from the database, but may remain on Cloudinary due to an error: {e}')

        return redirect('rfp_detail', pk=referral.pk)

    context = {
        'referral': referral,
        'rfp_file': rfp_file,
        'filename': filename,
    }
    return render(request, 'rfp/rfp_delete_file.html', context)

@login_required
def rfp_download_file(request, rfp_pk, file_pk):
    """
    Download file - now handles both HTML and DOCX with proper file type detection
    """
    referral = get_object_or_404(RFPReferral, pk=rfp_pk)
    rfp_file = get_object_or_404(RFPFile, pk=file_pk, referral=referral)
    user = request.user

    is_admin = hasattr(user, 'role') and user.role == 'admin'

    if not (is_admin or referral.client == user):
        return HttpResponseForbidden("You do not have permission to download this file.")

    # Check if user specifically wants DOCX format
    if request.GET.get('format') == 'docx':
        # Check if we have an original DOCX version
        if hasattr(rfp_file, 'original_docx') and rfp_file.original_docx and rfp_file.original_docx.cloudinary_url:
            # Download original DOCX
            print(f"Downloading original DOCX: {rfp_file.original_docx.cloudinary_url}")
            return HttpResponseRedirect(rfp_file.original_docx.cloudinary_url)
        else:
            # No original DOCX available, check if current file is DOCX
            is_current_docx = (
                not rfp_file.file_type == 'html' and
                not rfp_file.cloudinary_url.endswith('.html') and
                '_edited.html' not in rfp_file.cloudinary_url
            )
            if is_current_docx:
                # Current file is DOCX, download it
                print(f"Downloading current DOCX file: {rfp_file.cloudinary_url}")
                return HttpResponseRedirect(rfp_file.cloudinary_url)
            else:
                # Current file is HTML, show warning
                messages.warning(request,
                               'DOCX format not available for this file. The file has been edited online and is only available in HTML format for perfect formatting preservation.')
                return redirect('rfp_detail', pk=referral.pk)
    else:
        # Download current file (default - could be HTML or DOCX)
        print(f"Downloading current file: {rfp_file.cloudinary_url}")
        return HttpResponseRedirect(rfp_file.cloudinary_url)
def rfp_guide(request):
    """
    Display a comprehensive guide on how to submit RFPs to OB Global
    """
    return render(request, 'rfp/rfp_guide.html')


from docx import Document
import io
import tempfile
import os
from django.http import HttpResponse
import requests
from xml.etree import ElementTree
import zipfile


def html_to_docx_with_images(html_content):
    """
    Convert HTML to DOCX with image preservation
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from bs4 import BeautifulSoup
    import re
    import base64
    import io
    import tempfile
    import requests
    from PIL import Image

    doc = Document()

    # Set default styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    soup = BeautifulSoup(html_content, 'html.parser')

    def process_element(element, current_paragraph=None, inherited_styles=None):
        if inherited_styles is None:
            inherited_styles = {}

        # Handle text nodes
        if element.name is None:
            text = str(element).strip()
            if text and current_paragraph:
                run = current_paragraph.add_run(text)
                # Apply inherited styles
                if inherited_styles.get('bold'):
                    run.bold = True
                if inherited_styles.get('italic'):
                    run.italic = True
                if inherited_styles.get('underline'):
                    run.underline = True
            return current_paragraph

        # Get current element styles
        current_styles = inherited_styles.copy()

        # Extract styles from element
        style_attr = element.get('style', '')
        if 'font-weight: bold' in style_attr or element.name in ['strong', 'b']:
            current_styles['bold'] = True
        if 'font-style: italic' in style_attr or element.name in ['em', 'i']:
            current_styles['italic'] = True
        if 'text-decoration: underline' in style_attr or element.name == 'u':
            current_styles['underline'] = True

        # Handle different element types
        if element.name == 'p':
            current_paragraph = doc.add_paragraph()

            # Handle alignment
            if 'text-align: center' in style_attr:
                current_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif 'text-align: right' in style_attr:
                current_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif 'text-align: justify' in style_attr:
                current_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Process children
            for child in element.children:
                process_element(child, current_paragraph, current_styles)

        elif element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(element.name[1])
            heading = doc.add_heading(level=min(level - 1, 5))

            # Process heading content
            for child in element.children:
                process_element(child, heading, {'bold': True})  # Headings are typically bold

            current_paragraph = None

        elif element.name == 'br':
            if current_paragraph:
                current_paragraph.add_run().add_break()

        elif element.name == 'img':
            # Handle images - this is the key part!
            src = element.get('src', '')
            alt = element.get('alt', 'Image')

            if src:
                try:
                    # Handle base64 images
                    if src.startswith('data:image'):
                        # Extract base64 data
                        header, data = src.split(',', 1)
                        image_data = base64.b64decode(data)

                        # Create temporary file
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_file:
                            temp_file.write(image_data)
                            temp_path = temp_file.name

                        try:
                            # Add image to document
                            if current_paragraph is None:
                                current_paragraph = doc.add_paragraph()

                            # Calculate appropriate size
                            with Image.open(io.BytesIO(image_data)) as img:
                                width, height = img.size
                                # Scale down large images
                                max_width = 6.0  # inches
                                if width > 600:
                                    ratio = max_width / (width / 100)
                                    width_inches = max_width
                                else:
                                    width_inches = width / 100

                            current_paragraph.add_run().add_picture(temp_path, width=Inches(width_inches))

                        finally:
                            # Clean up temp file
                            if os.path.exists(temp_path):
                                os.unlink(temp_path)

                    # Handle external URL images
                    elif src.startswith('http'):
                        try:
                            response = requests.get(src, timeout=10)
                            if response.status_code == 200:
                                if current_paragraph is None:
                                    current_paragraph = doc.add_paragraph()

                                # Create temporary file
                                with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_file:
                                    temp_file.write(response.content)
                                    temp_path = temp_file.name

                                try:
                                    # Add image with reasonable size
                                    current_paragraph.add_run().add_picture(temp_path, width=Inches(4.0))
                                finally:
                                    if os.path.exists(temp_path):
                                        os.unlink(temp_path)
                        except Exception as e:
                            print(f"Failed to download image: {e}")
                            if current_paragraph:
                                current_paragraph.add_run(f"[Image: {alt}]").italic = True

                    else:
                        # Local or unknown image source
                        if current_paragraph:
                            current_paragraph.add_run(f"[Image: {alt}]").italic = True

                except Exception as e:
                    print(f"Error processing image: {e}")
                    if current_paragraph:
                        current_paragraph.add_run(f"[Image Error: {alt}]").italic = True

        elif element.name in ['ul', 'ol']:
            is_ordered = element.name == 'ol'
            for li in element.find_all('li', recursive=False):
                if is_ordered:
                    list_paragraph = doc.add_paragraph(style='List Number')
                else:
                    list_paragraph = doc.add_paragraph(style='List Bullet')

                for child in li.children:
                    process_element(child, list_paragraph, current_styles)

            current_paragraph = None

        elif element.name == 'table':
            # Handle tables
            rows = element.find_all('tr')
            if rows:
                # Determine number of columns
                col_count = max(len(row.find_all(['td', 'th'])) for row in rows)
                table = doc.add_table(rows=len(rows), cols=col_count)
                table.style = 'Table Grid'

                for i, row in enumerate(rows):
                    cells = row.find_all(['td', 'th'])
                    for j, cell in enumerate(cells):
                        if j < col_count:
                            # Clear default content and process cell content
                            cell_paragraph = table.cell(i, j).paragraphs[0]
                            for child in cell.children:
                                process_element(child, cell_paragraph, current_styles)

            current_paragraph = None

        else:
            # Process inline elements and unknown elements
            for child in element.children:
                current_paragraph = process_element(child, current_paragraph, current_styles)

        return current_paragraph

    # Start processing
    body = soup.find('body')
    if body:
        for child in body.children:
            process_element(child)
    else:
        for child in soup.children:
            process_element(child)

    # Ensure at least one paragraph exists
    if len(doc.paragraphs) == 0:
        doc.add_paragraph()

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    print(f"DOCX created with images, size: {len(buffer.getvalue())} bytes")
    return buffer.getvalue()


@login_required
def upload_image(request):
    """Handle image uploads for TinyMCE editor"""
    if request.method == 'POST' and request.FILES.get('file'):
        uploaded_file = request.FILES['file']

        # Validate image
        if not uploaded_file.content_type.startswith('image/'):
            return JsonResponse({'error': 'File must be an image'}, status=400)

        try:
            # Upload to Cloudinary
            upload_result = cloudinary.uploader.upload(
                uploaded_file,
                folder="rfp/editor-images",
                resource_type='image'
            )

            return JsonResponse({'location': upload_result['secure_url']})

        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)

    return JsonResponse({'error': 'Invalid request'}, status=400)

def extract_formatted_text_from_docx(docx_content):
    """
    Extract text from DOCX while preserving basic formatting
    """
    try:
        from docx import Document
        import io

        doc = Document(io.BytesIO(docx_content))
        html_content = []

        for paragraph in doc.paragraphs:
            para_html = []

            for run in paragraph.runs:
                text = run.text.strip()
                if text:
                    # Apply basic formatting
                    if run.bold:
                        text = f"<strong>{text}</strong>"
                    if run.italic:
                        text = f"<em>{text}</em>"
                    if run.underline:
                        text = f"<u>{text}</u>"

                    para_html.append(text)

            if para_html:
                html_content.append(f"<p>{' '.join(para_html)}</p>")

        # Handle tables
        for table in doc.tables:
            table_html = ['<table style="border: 1px solid black; border-collapse: collapse;">']
            for row in table.rows:
                table_html.append('<tr>')
                for cell in row.cells:
                    table_html.append(f'<td style="border: 1px solid black; padding: 5px;">{cell.text}</td>')
                table_html.append('</tr>')
            table_html.append('</table>')
            html_content.append(''.join(table_html))

        result = '\n'.join(html_content)
        return result if result else "<p>No readable content found.</p>"

    except Exception as e:
        print(f"Error extracting formatted text: {e}")
        return f"<p>Error extracting content: {str(e)}</p>"


def update_docx_with_html_content(original_docx_content, html_content):
    """
    Update DOCX document with HTML content while preserving original formatting
    """
    try:
        from docx import Document
        from bs4 import BeautifulSoup
        import io

        # Parse HTML content
        soup = BeautifulSoup(html_content, 'html.parser')

        # Load original document
        doc = Document(io.BytesIO(original_docx_content))

        # Clear existing content but preserve styles
        for element in doc.element.body:
            if element.tag.endswith('p') or element.tag.endswith('tbl'):
                doc.element.body.remove(element)

        # Process HTML and add to document
        def add_html_to_doc(element, current_paragraph=None):
            if element.name is None:
                # Text node
                if current_paragraph and element.strip():
                    run = current_paragraph.add_run(element.strip())
                return current_paragraph

            if element.name == 'p':
                current_paragraph = doc.add_paragraph()
                for child in element.children:
                    add_html_to_doc(child, current_paragraph)

            elif element.name in ['strong', 'b']:
                if current_paragraph:
                    run = current_paragraph.add_run(element.get_text())
                    run.bold = True

            elif element.name in ['em', 'i']:
                if current_paragraph:
                    run = current_paragraph.add_run(element.get_text())
                    run.italic = True

            elif element.name == 'u':
                if current_paragraph:
                    run = current_paragraph.add_run(element.get_text())
                    run.underline = True

            elif element.name == 'br':
                if current_paragraph:
                    current_paragraph.add_run().add_break()

            elif element.name in ['ul', 'ol']:
                for li in element.find_all('li', recursive=False):
                    if element.name == 'ul':
                        p = doc.add_paragraph(style='List Bullet')
                    else:
                        p = doc.add_paragraph(style='List Number')
                    add_html_to_doc(li, p)

            elif element.name == 'table':
                # Handle tables - create basic table structure
                rows = element.find_all('tr')
                if rows:
                    col_count = max(len(row.find_all(['td', 'th'])) for row in rows)
                    table = doc.add_table(rows=len(rows), cols=col_count)
                    table.style = 'Table Grid'

                    for i, row in enumerate(rows):
                        cells = row.find_all(['td', 'th'])
                        for j, cell in enumerate(cells):
                            if j < col_count:
                                table.cell(i, j).text = cell.get_text().strip()

            else:
                # Process children of unknown elements
                for child in element.children:
                    add_html_to_doc(child, current_paragraph)

            return current_paragraph

        # Process the HTML content
        body = soup.find('body')
        if body:
            for child in body.children:
                add_html_to_doc(child)
        else:
            for child in soup.children:
                add_html_to_doc(child)

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        print("DOCX updated successfully")
        return buffer.getvalue()

    except Exception as e:
        print(f"Error updating DOCX: {e}")
        return None


def create_simple_docx_from_html(html_content):
    """
    Create a simple DOCX from HTML content (fallback method)
    """
    from docx import Document
    from bs4 import BeautifulSoup
    import io

    doc = Document()
    soup = BeautifulSoup(html_content, 'html.parser')

    # Extract all text content
    all_text = soup.get_text()

    # Split into paragraphs
    paragraphs = [p.strip() for p in all_text.split('\n') if p.strip()]

    for para in paragraphs:
        doc.add_paragraph(para)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer.getvalue()
@login_required
def rfp_download_for_edit(request, rfp_pk, file_pk):
    """
    Download DOCX file for external editing
    """
    referral = get_object_or_404(RFPReferral, pk=rfp_pk)
    rfp_file = get_object_or_404(RFPFile, pk=file_pk, referral=referral)
    user = request.user

    if not (hasattr(user, 'role') and user.role == 'admin' or referral.client == user):
        return HttpResponseForbidden("You do not have permission to edit this file.")

    # Mark as being edited
    rfp_file.is_being_edited = True
    rfp_file.last_edited_by = user
    rfp_file.last_edited_at = timezone.now()
    rfp_file.save()

    messages.info(request,
                  'File marked for editing. After editing in Word, use the "Replace File" option to upload your changes.'
                  )
    return redirect('rfp_download_file', rfp_pk=rfp_pk, file_pk=file_pk)


@login_required
def rfp_finish_edit(request, rfp_pk, file_pk):
    """
    Mark file as no longer being edited
    """
    referral = get_object_or_404(RFPReferral, pk=rfp_pk)
    rfp_file = get_object_or_404(RFPFile, pk=file_pk, referral=referral)
    user = request.user

    if rfp_file.last_edited_by == user or hasattr(user, 'role') and user.role == 'admin':
        rfp_file.is_being_edited = False
        rfp_file.save()
        messages.success(request, 'Editing session completed.')

    return redirect('rfp_detail', pk=referral.pk)


@login_required
def rfp_replace_file(request, rfp_pk, file_pk):
    """
    Replace an existing file with a new version
    """
    referral = get_object_or_404(RFPReferral, pk=rfp_pk)
    rfp_file = get_object_or_404(RFPFile, pk=file_pk, referral=referral)
    user = request.user

    is_admin = hasattr(user, 'role') and user.role == 'admin'
    if not (is_admin or referral.client == user):
        return HttpResponseForbidden("You do not have permission to replace this file.")

    if request.method == 'POST':
        uploaded_file = request.FILES.get('replacement_file')

        if not uploaded_file:
            messages.error(request, 'Please select a file to upload.')
            return redirect('rfp_detail', pk=referral.pk)

        # Check file extension
        file_name = uploaded_file.name.lower()
        if not (file_name.endswith('.docx') or file_name.endswith('.doc')):
            messages.error(request, 'Please upload a DOC or DOCX file.')
            return redirect('rfp_detail', pk=referral.pk)

        try:
            # STEP 1: Delete existing files from Cloudinary
            print(f"=== STARTING FILE REPLACEMENT PROCESS ===")
            print(f"1. Deleting existing files from Cloudinary...")

            # Get current file URLs before replacement
            current_file_url = rfp_file.cloudinary_url
            current_public_id = rfp_file.public_id

            print(f"1a. Current file URL: {current_file_url}")
            print(f"1b. Current public_id: {current_public_id}")

            # Check if there's an OriginalDocxFile record
            original_docx_exists = hasattr(rfp_file, 'original_docx')
            original_docx_url = None
            original_docx_public_id = None

            if original_docx_exists:
                original_docx_url = rfp_file.original_docx.cloudinary_url
                # Extract public_id from original docx URL
                if original_docx_url:
                    try:
                        # Extract public_id from Cloudinary URL
                        import re
                        match = re.search(r'/upload/(?:v\d+/)?(.+?)(?:\.[^.]*)?$', original_docx_url)
                        if match:
                            original_docx_public_id = match.group(1)
                    except Exception as e:
                        print(f"1c. Could not extract public_id from original docx: {e}")

            print(f"1d. Original DOCX exists: {original_docx_exists}")
            print(f"1e. Original DOCX URL: {original_docx_url}")
            print(f"1f. Original DOCX public_id: {original_docx_public_id}")

            # Delete current file from Cloudinary
            try:
                if current_public_id:
                    # Delete the main file
                    cloudinary.uploader.destroy(current_public_id, resource_type='raw')
                    print(f"1g. Deleted main file: {current_public_id}")

                    # Also try to delete the _edited version if it exists
                    edited_public_id = f"{current_public_id}_edited"
                    cloudinary.uploader.destroy(edited_public_id, resource_type='raw')
                    print(f"1h. Deleted edited version: {edited_public_id}")
            except Exception as e:
                print(f"1i. Warning: Could not delete main file from Cloudinary: {e}")

            # Delete original DOCX file from Cloudinary if it exists and is different from current file
            try:
                if original_docx_exists and original_docx_public_id and original_docx_public_id != current_public_id:
                    cloudinary.uploader.destroy(original_docx_public_id, resource_type='raw')
                    print(f"1j. Deleted original DOCX file: {original_docx_public_id}")
            except Exception as e:
                print(f"1k. Warning: Could not delete original DOCX file from Cloudinary: {e}")

            # STEP 2: Delete OriginalDocxFile record from database
            if original_docx_exists:
                try:
                    original_docx_record = rfp_file.original_docx
                    original_docx_record.delete()
                    print(f"2. Deleted OriginalDocxFile record from database")
                except Exception as e:
                    print(f"2a. Warning: Could not delete OriginalDocxFile record: {e}")

            # STEP 3: Upload new file to Cloudinary
            print(f"3. Uploading new file to Cloudinary...")
            upload_result = cloudinary.uploader.upload(
                uploaded_file,
                public_id=current_public_id,  # Use the same public_id to maintain consistency
                resource_type='raw',
                overwrite=True,
                access_mode='public'
            )

            print(f"3a. New file uploaded successfully: {upload_result.get('secure_url')}")

            # STEP 4: Update RFPFile record
            print(f"4. Updating RFPFile record...")
            rfp_file.cloudinary_url = upload_result.get('secure_url')
            rfp_file.file_type = 'docx'  # Reset to docx since we're uploading a new DOCX file
            rfp_file.is_being_edited = False
            rfp_file.last_edited_by = None
            rfp_file.last_edited_at = None
            rfp_file.uploaded_by = user
            rfp_file.uploaded_at = timezone.now()
            rfp_file.save()

            print(f"4a. RFPFile updated with new URL: {rfp_file.cloudinary_url}")

            # STEP 5: Send notification
            print(f"5. Sending notification...")
            send_rfp_update_notification(referral, user, action='file_replace')

            print(f"=== FILE REPLACEMENT COMPLETED ===")
            messages.success(request, 'File replaced successfully! All previous versions have been cleaned up.')

        except Exception as e:
            error_msg = f'Error replacing file: {str(e)}'
            print(f"=== FILE REPLACEMENT FAILED ===")
            print(f"ERROR: {error_msg}")
            import traceback
            print(f"TRACEBACK: {traceback.format_exc()}")
            messages.error(request, error_msg)

    return redirect('rfp_detail', pk=referral.pk)
# rfp/docx_utils.py
import zipfile
import xml.etree.ElementTree as ET
import io


def extract_text_from_docx(file_content):
    """
    Extract text content from DOCX file using built-in Python libraries
    """
    try:
        # Create a zip file from the bytes
        with zipfile.ZipFile(io.BytesIO(file_content), 'r') as docx_zip:
            # Read the main document XML
            if 'word/document.xml' in docx_zip.namelist():
                with docx_zip.open('word/document.xml') as document_file:
                    xml_content = document_file.read()

                    # Parse XML and extract text
                    root = ET.fromstring(xml_content)

                    # Namespace for DOCX
                    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

                    # Extract text from all paragraphs
                    paragraphs = []
                    for paragraph in root.findall('.//w:p', ns):
                        text_elements = []
                        for text_elem in paragraph.findall('.//w:t', ns):
                            if text_elem.text:
                                text_elements.append(text_elem.text)

                        if text_elements:
                            paragraph_text = ''.join(text_elements).strip()
                            if paragraph_text:
                                paragraphs.append(paragraph_text)

                    extracted_text = '\n'.join(paragraphs)
                    print(f"Extracted {len(paragraphs)} paragraphs, {len(extracted_text)} characters")
                    return extracted_text if extracted_text else "No text content found in document."

        return "Could not find document content in the file."

    except Exception as e:
        error_msg = f"Error extracting text: {str(e)}"
        print(error_msg)
        return error_msg

@login_required
def rfp_edit_docx_enhanced(request, rfp_pk, file_pk):
    """
    Enhanced DOCX editor with TinyMCE and mammoth for format preservation
    """
    import mammoth

    referral = get_object_or_404(RFPReferral, pk=rfp_pk)
    rfp_file = get_object_or_404(RFPFile, pk=file_pk, referral=referral)
    user = request.user

    is_admin = hasattr(user, 'role') and user.role == 'admin'

    # Permission check
    if not (is_admin or referral.client == user):
        return HttpResponseForbidden("You do not have permission to edit this file.")

    # Initialize text_content
    text_content = ""

    try:
        # Download the file from Cloudinary
        print(f"Downloading file from: {rfp_file.cloudinary_url}")
        response = requests.get(rfp_file.cloudinary_url, timeout=30)
        response.raise_for_status()

        # Convert DOCX to HTML using mammoth (preserves formatting)
        print("Converting DOCX to HTML using mammoth...")
        docx_file = io.BytesIO(response.content)
        result = mammoth.convert_to_html(docx_file)
        text_content = result.value

        if result.messages:
            print(f"Mammoth conversion messages: {result.messages}")

        print(f"Converted HTML length: {len(text_content)}")

        if not text_content or text_content.strip() == "":
            text_content = "<p>No readable content found. You can start typing below.</p>"
        else:
            print("DOCX successfully converted to HTML with formatting preserved")

    except Exception as e:
        error_msg = f"Could not load document: {str(e)}"
        print(error_msg)
        text_content = f"<p>Unable to load existing document content. You can start typing below.</p><p><strong>Error:</strong> {str(e)}</p>"
        messages.warning(request, "Could not extract content from the document.")

    if request.method == 'POST':
        try:
            edited_content = request.POST.get('docx_content', '')
            is_ajax = request.headers.get('X-Requested-With') == 'XMLHttpRequest'

            if not edited_content.strip():
                if is_ajax:
                    return JsonResponse({'status': 'error', 'message': 'Document content cannot be empty.'})
                messages.error(request, "Document content cannot be empty.")
                context = {
                    'referral': referral,
                    'rfp_file': rfp_file,
                    'text_content': edited_content,
                    'is_admin': is_admin,
                }
                return render(request, 'rfp/rfp_edit_docx_tinymce.html', context)

            # Convert HTML to DOCX using python-docx with comprehensive parsing
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from bs4 import BeautifulSoup
            import re

            doc = Document()

            # Set default font
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(12)

            soup = BeautifulSoup(edited_content, 'html.parser')

            def hex_to_rgb(hex_color):
                hex_color = hex_color.lstrip('#')
                if len(hex_color) == 6:
                    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
                return None

            def extract_styles(element):
                styles = {}
                style_attr = element.get('style', '')
                if 'font-weight: bold' in style_attr or element.name in ['strong', 'b']:
                    styles['bold'] = True
                if 'font-style: italic' in style_attr or element.name in ['em', 'i']:
                    styles['italic'] = True
                if 'text-decoration: underline' in style_attr or element.name == 'u':
                    styles['underline'] = True
                return styles

            # Process HTML elements
            for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'ul', 'ol', 'table']):
                if element.name in ['h1', 'h2', 'h3']:
                    p = doc.add_heading(element.get_text().strip(), level=int(element.name[1]))
                elif element.name == 'p':
                    p = doc.add_paragraph()
                    for child in element.descendants:
                        if isinstance(child, str) and child.strip():
                            run = p.add_run(child)
                            parent = child.parent
                            styles = extract_styles(parent)
                            if styles.get('bold'):
                                run.bold = True
                            if styles.get('italic'):
                                run.italic = True
                            if styles.get('underline'):
                                run.underline = True
                elif element.name in ['ul', 'ol']:
                    for li in element.find_all('li', recursive=False):
                        p = doc.add_paragraph(li.get_text().strip(), style='List Bullet' if element.name == 'ul' else 'List Number')
                elif element.name == 'table':
                    rows = element.find_all('tr')
                    if rows:
                        cols = max(len(row.find_all(['td', 'th'])) for row in rows)
                        table = doc.add_table(rows=len(rows), cols=cols)
                        for i, row in enumerate(rows):
                            cells = row.find_all(['td', 'th'])
                            for j, cell in enumerate(cells):
                                table.rows[i].cells[j].text = cell.get_text().strip()

            # Save to buffer
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)

            # Upload to Cloudinary
            upload_result = cloudinary.uploader.upload(
                doc_buffer,
                public_id=rfp_file.public_id,
                resource_type='raw',
                overwrite=True,
                format='docx'
            )

            rfp_file.cloudinary_url = upload_result['secure_url']
            rfp_file.save()

            # Return JSON for AJAX requests (auto-save)
            if is_ajax:
                return JsonResponse({'status': 'success', 'message': 'Document saved successfully'})

            # Regular form submission
            messages.success(request, f"Document '{rfp_file.get_file_name}' has been updated successfully!")
            return redirect('rfp_detail', pk=referral.pk)

        except Exception as e:
            error_msg = f"Error saving document: {str(e)}"
            print(error_msg)
            if is_ajax:
                return JsonResponse({'status': 'error', 'message': str(e)})
            messages.error(request, f"Failed to save document: {str(e)}")

    context = {
        'referral': referral,
        'rfp_file': rfp_file,
        'text_content': text_content,
        'is_admin': is_admin,
    }
    return render(request, 'rfp/rfp_edit_docx_tinymce.html', context)



def enhanced_html_to_docx(html_content):
    """
    Enhanced HTML to DOCX conversion with maximum format preservation
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml.ns import qn
    from bs4 import BeautifulSoup
    import re

    doc = Document()

    # Set default styles to match typical Word document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Remove the empty default paragraph
    if len(doc.paragraphs) > 0:
        doc._body._body.remove(doc.paragraphs[0]._p)

    soup = BeautifulSoup(html_content, 'html.parser')

    def hex_to_rgb(hex_color):
        """Convert hex color to RGB"""
        hex_color = hex_color.lstrip('#')
        if len(hex_color) == 6:
            return tuple(int(hex_color[i:i + 2], 16) for i in (0, 2, 4))
        return None

    def extract_styles_from_element(element):
        """Extract comprehensive styling from HTML element"""
        styles = {}
        style_attr = element.get('style', '')

        # Text formatting
        if 'font-weight: bold' in style_attr or element.name in ['strong', 'b']:
            styles['bold'] = True
        if 'font-style: italic' in style_attr or element.name in ['em', 'i']:
            styles['italic'] = True
        if 'text-decoration: underline' in style_attr or element.name == 'u':
            styles['underline'] = True
        if 'text-decoration: line-through' in style_attr or element.name in ['s', 'strike']:
            styles['strike'] = True

        # Font properties
        font_family_match = re.search(r'font-family:\s*([^;]+)', style_attr)
        if font_family_match:
            styles['font_family'] = font_family_match.group(1).strip().strip('"\'')

        font_size_match = re.search(r'font-size:\s*(\d+)(?:px|pt)', style_attr)
        if font_size_match:
            styles['font_size'] = int(font_size_match.group(1))

        # Colors
        color_match = re.search(r'color:\s*(#[0-9a-fA-F]+|[a-zA-Z]+)', style_attr)
        if color_match:
            color = color_match.group(1)
            if color.startswith('#'):
                rgb = hex_to_rgb(color)
                if rgb:
                    styles['color'] = RGBColor(*rgb)

        # Background color
        bg_match = re.search(r'background-color:\s*(#[0-9a-fA-F]+|[a-zA-Z]+)', style_attr)
        if bg_match:
            bg_color = bg_match.group(1)
            if bg_color.startswith('#'):
                rgb = hex_to_rgb(bg_color)
                if rgb:
                    styles['bg_color'] = RGBColor(*rgb)

        # Text alignment
        if 'text-align: center' in style_attr:
            styles['alignment'] = WD_ALIGN_PARAGRAPH.CENTER
        elif 'text-align: right' in style_attr:
            styles['alignment'] = WD_ALIGN_PARAGRAPH.RIGHT
        elif 'text-align: justify' in style_attr:
            styles['alignment'] = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif 'text-align: left' in style_attr:
            styles['alignment'] = WD_ALIGN_PARAGRAPH.LEFT

        return styles

    def apply_styles_to_run(run, styles):
        """Apply styles to a docx run"""
        if styles.get('bold'):
            run.bold = True
        if styles.get('italic'):
            run.italic = True
        if styles.get('underline'):
            run.underline = True
        if styles.get('strike'):
            run.font.strike = True
        if styles.get('font_family'):
            run.font.name = styles['font_family']
        if styles.get('font_size'):
            run.font.size = Pt(styles['font_size'])
        if styles.get('color'):
            run.font.color.rgb = styles['color']

    def apply_styles_to_paragraph(paragraph, styles):
        """Apply styles to a docx paragraph"""
        if styles.get('alignment'):
            paragraph.alignment = styles['alignment']

    def process_element(element, current_paragraph=None, inherited_styles=None):
        """Recursively process HTML elements and convert to DOCX"""
        if inherited_styles is None:
            inherited_styles = {}

        # Handle text nodes
        if element.name is None:
            text = str(element).strip()
            if text and current_paragraph:
                # Add text to current paragraph
                run = current_paragraph.add_run(text)
                apply_styles_to_run(run, inherited_styles)
            return current_paragraph

        # Get styles for current element
        current_styles = extract_styles_from_element(element)
        # Merge with inherited styles
        combined_styles = {**inherited_styles, **current_styles}

        if element.name == 'p':
            # Create new paragraph for each <p> tag
            current_paragraph = doc.add_paragraph()
            apply_styles_to_paragraph(current_paragraph, combined_styles)

            # Process all children
            for child in element.children:
                current_paragraph = process_element(child, current_paragraph, combined_styles)

        elif element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            # Handle headings
            level = int(element.name[1])
            heading = doc.add_heading(level=min(level - 1, 5))  # docx supports levels 0-5
            apply_styles_to_paragraph(heading, combined_styles)

            # Process heading content
            for child in element.children:
                process_element(child, heading, combined_styles)
            current_paragraph = None

        elif element.name == 'br':
            # Line break
            if current_paragraph:
                current_paragraph.add_run().add_break(WD_BREAK.LINE)

        elif element.name in ['strong', 'b', 'em', 'i', 'u', 'span', 's', 'strike']:
            # Handle inline formatting - continue in current paragraph
            for child in element.children:
                current_paragraph = process_element(child, current_paragraph, combined_styles)

        elif element.name == 'table':
            # Handle tables with better preservation
            rows = element.find_all('tr')
            if not rows:
                return current_paragraph

            # Count columns from the first row
            first_row = rows[0]
            cols = len(first_row.find_all(['td', 'th']))

            # Create table with proper dimensions
            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = 'Table Grid'

            # Process each row
            for row_idx, tr in enumerate(rows):
                if row_idx >= len(table.rows):
                    break

                row_cells = tr.find_all(['td', 'th'])
                table_row = table.rows[row_idx]

                # Process each cell in the row
                for col_idx, cell in enumerate(row_cells):
                    if col_idx >= len(table_row.cells):
                        break

                    cell_paragraph = table_row.cells[col_idx].paragraphs[0]
                    # Clear default content
                    cell_paragraph.clear()

                    # Process cell content
                    for child in cell.children:
                        process_element(child, cell_paragraph, combined_styles)

            current_paragraph = None

        elif element.name in ['ul', 'ol']:
            # Handle lists
            is_ordered = element.name == 'ol'
            for li in element.find_all('li'):
                if is_ordered:
                    list_paragraph = doc.add_paragraph(style='List Number')
                else:
                    list_paragraph = doc.add_paragraph(style='List Bullet')

                # Process list item content
                for child in li.children:
                    process_element(child, list_paragraph, combined_styles)
            current_paragraph = None

        elif element.name == 'img':
            # Handle images - add placeholder text
            alt_text = element.get('alt', 'Image')
            if current_paragraph:
                run = current_paragraph.add_run(f"[Image: {alt_text}]")
                run.italic = True
                run.font.color.rgb = RGBColor(128, 128, 128)

        elif element.name == 'a':
            # Handle links
            href = element.get('href', '#')
            link_text = element.get_text()
            if current_paragraph:
                run = current_paragraph.add_run(link_text)
                run.font.color.rgb = RGBColor(0, 0, 255)
                run.underline = True
                # Note: python-docx doesn't support hyperlinks in basic usage

        elif element.name == 'blockquote':
            # Handle blockquotes with indentation
            current_paragraph = doc.add_paragraph()
            current_paragraph.paragraph_format.left_indent = Inches(0.5)
            current_paragraph.paragraph_format.right_indent = Inches(0.5)

            for child in element.children:
                current_paragraph = process_element(child, current_paragraph, combined_styles)

        elif element.name == 'code':
            # Handle code blocks
            if current_paragraph:
                run = current_paragraph.add_run(element.get_text())
                run.font.name = 'Courier New'
                run.font.size = Pt(10)

        elif element.name == 'div':
            # Handle divs - process children
            for child in element.children:
                current_paragraph = process_element(child, current_paragraph, combined_styles)

        else:
            # Process all children for unknown elements
            for child in element.children:
                current_paragraph = process_element(child, current_paragraph, combined_styles)

        return current_paragraph

    # Start processing from the body or root
    body = soup.find('body')
    if body:
        # If there's a body tag, process its children
        for child in body.children:
            process_element(child)
    else:
        # Otherwise process the root children
        for child in soup.children:
            process_element(child)

    # If no content was added (empty editor), add a blank paragraph
    if len(doc.paragraphs) == 0:
        doc.add_paragraph()

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    print(f"DOCX file created successfully. Size: {len(buffer.getvalue())} bytes")
    return buffer


@login_required
def rfp_edit_docx(request, rfp_pk, file_pk):
    """
    Edit document content and save as HTML for perfect preservation
    """
    import mammoth
    import io
    import base64
    import requests

    referral = get_object_or_404(RFPReferral, pk=rfp_pk)
    rfp_file = get_object_or_404(RFPFile, pk=file_pk, referral=referral)
    user = request.user

    is_admin = hasattr(user, 'role') and user.role == 'admin'

    if not (is_admin or referral.client == user):
        return HttpResponseForbidden("You do not have permission to edit this file.")

    # Initialize text_content
    text_content = ""

    try:
        # Download the file from Cloudinary
        print(f"Downloading file from: {rfp_file.cloudinary_url}")
        response = requests.get(rfp_file.cloudinary_url, timeout=30)
        response.raise_for_status()

        # CHECK FILE TYPE - If it's already HTML, just read the content directly
        is_html_file = (
                rfp_file.file_type == 'html' or
                rfp_file.cloudinary_url.endswith('.html') or
                '_edited.html' in rfp_file.cloudinary_url
        )

        if is_html_file:
            print("File is already HTML - reading content directly")
            # It's an HTML file, read it directly
            text_content = response.text

            # Extract just the body content from the full HTML document
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(text_content, 'html.parser')
            body_content = soup.find('body')
            if body_content:
                # Get the content inside the document-container div or just the body
                container = body_content.find('div', class_='document-container')
                if container:
                    text_content = str(container)
                else:
                    text_content = str(body_content)
            else:
                # If no body found, try to extract content from document-container directly
                container = soup.find('div', class_='document-container')
                if container:
                    text_content = str(container)

            print(f"Direct HTML content length: {len(text_content)}")

        else:
            # It's a DOCX file, convert using mammoth
            print("Converting DOCX to HTML using mammoth with images...")
            docx_file = io.BytesIO(response.content)

            # Convert with image support
            def convert_image(image):
                with image.open() as image_bytes:
                    image_data = image_bytes.read()
                    encoded_src = base64.b64encode(image_data).decode("ascii")

                    # Determine image type
                    if image.content_type == "image/png":
                        mime_type = "png"
                    elif image.content_type == "image/jpeg":
                        mime_type = "jpeg"
                    elif image.content_type == "image/gif":
                        mime_type = "gif"
                    else:
                        mime_type = "png"

                    return {"src": f"data:image/{mime_type};base64,{encoded_src}"}

            result = mammoth.convert_to_html(
                docx_file,
                convert_image=mammoth.images.img_element(convert_image)
            )

            text_content = result.value

            if result.messages:
                print(f"Mammoth conversion messages: {result.messages}")

            print(f"Converted HTML length: {len(text_content)}")

        if not text_content or text_content.strip() == "":
            text_content = "<p>No readable content found. You can start typing below.</p>"
            print("No content found in document")
        else:
            print("Document loaded successfully")

    except Exception as e:
        error_msg = f"Could not load document: {str(e)}"
        print(error_msg)
        text_content = f"<p>Unable to load existing document content. You can start typing below.</p><p><strong>Error:</strong> {str(e)}</p>"
        messages.warning(request, "Could not extract content from the document.")

    if request.method == 'POST':
        try:
            edited_content = request.POST.get('docx_content', '')
            is_ajax = request.headers.get('X-Requested-With') == 'XMLHttpRequest'

            print(f"=== STARTING SAVE PROCESS ===")
            print(f"Received content for saving, length: {len(edited_content)}")
            print(f"Is AJAX request: {is_ajax}")

            if not edited_content.strip():
                if is_ajax:
                    return JsonResponse({'status': 'error', 'message': 'Document content cannot be empty.'})
                messages.error(request, "Document content cannot be empty.")
                context = {
                    'referral': referral,
                    'rfp_file': rfp_file,
                    'text_content': edited_content,
                    'is_admin': is_admin,
                }
                return render(request, 'rfp/rfp_edit_docx_tinymce.html', context)

            # Save the current file URL before we change it
            current_file_url = rfp_file.cloudinary_url
            print(f"1. Current file URL: {current_file_url}")
            print(f"2. Current public_id: {rfp_file.public_id}")

            # Create enhanced HTML file with proper styling
            enhanced_html = create_enhanced_html(edited_content, rfp_file.get_file_name())
            print(f"3. Enhanced HTML created, length: {len(enhanced_html)}")

            # Upload HTML to Cloudinary
            print("4. Starting HTML upload to Cloudinary...")
            html_file = io.BytesIO(enhanced_html.encode('utf-8'))

            try:
                # FIXED: Add access_mode='public' and proper folder structure like rfp_add_file
                upload_result = cloudinary.uploader.upload(
                    html_file,
                    public_id=f"{rfp_file.public_id}_edited",
                    resource_type='raw',
                    access_mode='public',  # ADD THIS - CRITICAL FOR PUBLIC ACCESS
                    overwrite=True,
                    format='html',
                    timeout=600  # ADD THIS - match timeout from rfp_add_file
                )
                print(f"4b. HTML upload successful!")
                print(f"4c. New HTML URL: {upload_result.get('secure_url')}")

                # Verify the upload was successful and publicly accessible
                test_response = requests.head(upload_result['secure_url'], timeout=10)
                if test_response.status_code == 200:
                    print("4d. HTML file is publicly accessible")
                else:
                    print(f"4d. WARNING: HTML file may not be publicly accessible: {test_response.status_code}")

            except Exception as upload_error:
                print(f"4f. HTML upload FAILED: {upload_error}")
                raise upload_error

            # Update file record to point to HTML version
            print("5. Updating RFPFile record...")
            rfp_file.cloudinary_url = upload_result['secure_url']
            rfp_file.file_type = 'html'
            rfp_file.save()
            print(f"6. RFPFile updated with HTML URL: {rfp_file.cloudinary_url}")

            # FIXED: SIMPLIFIED ORIGINAL DOCX CREATION LOGIC
            # Always create/update OriginalDocxFile if we don't have one OR if we're converting from DOCX
            should_create_original_docx = False

            # Check if we already have an OriginalDocxFile
            has_original_docx = hasattr(rfp_file, 'original_docx') and rfp_file.original_docx.cloudinary_url

            # Determine if current file is a DOCX file (before conversion)
            is_original_file_docx = (
                    rfp_file.file_type == 'docx' or
                    current_file_url.lower().endswith(('.docx', '.doc')) or
                    (not current_file_url.endswith('.html') and '_edited.html' not in current_file_url)
            )

            print(f"7. Checking OriginalDocxFile creation:")
            print(f"7a. Has original docx: {has_original_docx}")
            print(f"7b. Is original file docx: {is_original_file_docx}")
            print(f"7c. Current file URL: {current_file_url}")

            if not has_original_docx and is_original_file_docx:
                print("7d. Creating OriginalDocxFile...")
                try:
                    # Verify the original DOCX is accessible
                    original_test_response = requests.head(current_file_url, timeout=10)
                    if original_test_response.status_code == 200:
                        print("7e. Original DOCX is publicly accessible")

                        # Create or update OriginalDocxFile
                        original_docx, created = OriginalDocxFile.objects.get_or_create(
                            rfp_file=rfp_file,
                            defaults={'cloudinary_url': current_file_url}
                        )

                        if not created:
                            # Update existing record
                            original_docx.cloudinary_url = current_file_url
                            original_docx.save()

                        print(
                            f"8. OriginalDocxFile {'created' if created else 'updated'}: {original_docx.cloudinary_url}")
                    else:
                        print(
                            f"7e. WARNING: Original DOCX may not be publicly accessible: {original_test_response.status_code}")
                        # Still create the record but log the warning
                        original_docx, created = OriginalDocxFile.objects.get_or_create(
                            rfp_file=rfp_file,
                            defaults={'cloudinary_url': current_file_url}
                        )
                        print(f"8. OriginalDocxFile created with accessibility warning")

                except Exception as e:
                    print(f"8a. Error creating/updating OriginalDocxFile: {e}")
                    # Don't fail the entire process if OriginalDocxFile creation fails
                    import traceback
                    print(f"8b. Traceback: {traceback.format_exc()}")
            elif has_original_docx:
                print("7d. OriginalDocxFile already exists, no need to create")
            else:
                print("7d. Not creating OriginalDocxFile - file is not DOCX or already HTML")

            # Send notification
            print("9. Sending notification...")
            send_rfp_update_notification(referral, request.user, action='file_edit')

            print("10. HTML document saved successfully!")
            print("=== SAVE PROCESS COMPLETED ===")

            if is_ajax:
                return JsonResponse({
                    'status': 'success',
                    'message': 'Document saved successfully! All images and formatting preserved perfectly.'
                })

            messages.success(request, 'Document updated successfully! All images and formatting preserved.')
            return redirect('rfp_detail', pk=referral.pk)

        except Exception as e:
            error_msg = f'Error saving document: {str(e)}'
            print(f"=== SAVE PROCESS FAILED ===")
            print(f"ERROR: {error_msg}")
            import traceback
            print(f"TRACEBACK: {traceback.format_exc()}")
            print("=== END ERROR ===")

            if is_ajax:
                return JsonResponse({
                    'status': 'error',
                    'message': f'Failed to save document: {str(e)}'
                })

            messages.error(request, error_msg)
            context = {
                'referral': referral,
                'rfp_file': rfp_file,
                'text_content': request.POST.get('docx_content', ''),
                'is_admin': is_admin,
            }
            return render(request, 'rfp/rfp_edit_docx_tinymce.html', context)

    context = {
        'referral': referral,
        'rfp_file': rfp_file,
        'text_content': text_content,
        'is_admin': is_admin,
        'debug': True
    }
    return render(request, 'rfp/rfp_edit_docx_tinymce.html', context)
def create_enhanced_html(html_content, filename):
    """Create a well-formatted HTML file with proper styling"""
    enhanced_html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{filename}</title>
    <style>
        body {{
            font-family: 'Times New Roman', Times, serif;
            font-size: 12pt;
            line-height: 1.6;
            margin: 2cm;
            color: #000000;
            background: #ffffff;
        }}
        .document-container {{
            max-width: 21cm;
            margin: 0 auto;
            padding: 2cm;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 1em 0;
        }}
        table, th, td {{
            border: 1px solid #333;
        }}
        th, td {{
            padding: 8px 12px;
            text-align: left;
        }}
        th {{
            background-color: #f2f2f2;
            font-weight: bold;
        }}
        img {{
            max-width: 100%;
            height: auto;
            display: block;
            margin: 1em auto;
        }}
        p {{
            margin: 0 0 1em 0;
        }}
        h1, h2, h3, h4, h5, h6 {{
            margin: 1.5em 0 0.5em 0;
            color: #000000;
        }}
        h1 {{ font-size: 18pt; }}
        h2 {{ font-size: 16pt; }}
        h3 {{ font-size: 14pt; }}
        ul, ol {{
            margin: 1em 0;
            padding-left: 2em;
        }}
        li {{
            margin: 0.5em 0;
        }}
        .page-break {{
            page-break-after: always;
        }}
        @media print {{
            body {{
                margin: 0;
                padding: 0;
            }}
            .document-container {{
                padding: 0;
                margin: 0;
            }}
        }}
    </style>
</head>
<body>
    <div class="document-container">
        {html_content}
    </div>
</body>
</html>"""
    return enhanced_html


@login_required
def rfp_download_file(request, rfp_pk, file_pk):
    """
    Download file - now handles both HTML and DOCX
    """
    referral = get_object_or_404(RFPReferral, pk=rfp_pk)
    rfp_file = get_object_or_404(RFPFile, pk=file_pk, referral=referral)
    user = request.user

    is_admin = hasattr(user, 'role') and user.role == 'admin'

    if not (is_admin or referral.client == user):
        return HttpResponseForbidden("You do not have permission to download this file.")

    # Check if we have an HTML version and original DOCX reference
    if hasattr(rfp_file, 'original_docx') and rfp_file.original_docx.cloudinary_url:
        # Offer choice between HTML (perfect) and original DOCX
        if request.GET.get('format') == 'docx':
            # Download original DOCX
            print(f"Downloading original DOCX: {rfp_file.original_docx.cloudinary_url}")
            return HttpResponseRedirect(rfp_file.original_docx.cloudinary_url)
        else:
            # Download HTML version (default - perfect preservation)
            print(f"Downloading HTML version: {rfp_file.cloudinary_url}")
            return HttpResponseRedirect(rfp_file.cloudinary_url)
    else:
        # Original DOCX only (no HTML version available yet)
        print(f"Downloading original file: {rfp_file.cloudinary_url}")
        return HttpResponseRedirect(rfp_file.cloudinary_url)
    

from django.core.cache import cache


def get_cached_conversion(rfp_file):
    cache_key = f"docx_conversion_{rfp_file.pk}_{rfp_file.updated_at.timestamp()}"
    cached_content = cache.get(cache_key)

    if cached_content:
        return cached_content

    # Perform conversion
    converted_content = convert_docx_to_html(rfp_file)

    # Cache for 1 hour
    cache.set(cache_key, converted_content, 3600)
    return converted_content


def sanitize_html(html_content):
    """Basic HTML sanitization"""
    allowed_tags = {'p', 'br', 'strong', 'em', 'u', 'ul', 'ol', 'li', 'h1', 'h2', 'h3', 'h4', 'table', 'tr', 'td', 'th',
                    'img'}
    allowed_attrs = {'src', 'alt', 'title', 'style', 'class', 'border', 'cellpadding', 'cellspacing'}

    # Use a proper HTML sanitizer like bleach in production
    import bleach
    return bleach.clean(html_content, tags=allowed_tags, attributes=allowed_attrs)


def convert_image(image):
    try:
        with image.open() as image_bytes:
            image_data = image_bytes.read()

            # Validate image size
            if len(image_data) > 10 * 1024 * 1024:  # 10MB
                return {"src": "#"}

            encoded_src = base64.b64encode(image_data).decode("ascii")

            # Extended MIME type detection
            mime_map = {
                "image/png": "png",
                "image/jpeg": "jpeg",
                "image/gif": "gif",
                "image/webp": "webp",
                "image/svg+xml": "svg+xml"
            }
            mime_type = mime_map.get(image.content_type, "png")

            return {"src": f"data:image/{mime_type};base64,{encoded_src}"}
    except Exception as e:
        logger.warning(f"Failed to process image: {e}")
        return {"src": "#"}


@login_required
def rfp_download_original_docx(request, rfp_pk, file_pk):
    """
    Download the original DOCX file for an edited HTML file
    """
    referral = get_object_or_404(RFPReferral, pk=rfp_pk)
    rfp_file = get_object_or_404(RFPFile, pk=file_pk, referral=referral)

    # Check permissions
    is_admin = hasattr(request.user, 'role') and request.user.role == 'admin'
    if not (is_admin or referral.client == request.user):
        return HttpResponseForbidden("You do not have permission to download this file.")

    # Check if original DOCX exists
    if not hasattr(rfp_file, 'original_docx') or not rfp_file.original_docx.cloudinary_url:
        messages.error(request, "Original DOCX file not found.")
        return redirect('rfp_detail', pk=referral.pk)

    try:
        # Download the original DOCX from Cloudinary
        response = requests.get(rfp_file.original_docx.cloudinary_url, timeout=30)
        response.raise_for_status()

        # Create response with original DOCX content
        file_response = HttpResponse(response.content,
                                     content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

        # Set filename
        original_filename = rfp_file.get_file_name()
        if not original_filename.lower().endswith(('.docx', '.doc')):
            original_filename += '.docx'

        file_response['Content-Disposition'] = f'attachment; filename="{original_filename}"'
        return file_response

    except Exception as e:
        error_msg = f"Could not download original DOCX file: {str(e)}"
        print(error_msg)
        messages.error(request, "Failed to download original DOCX file.")
        return redirect('rfp_detail', pk=referral.pk)